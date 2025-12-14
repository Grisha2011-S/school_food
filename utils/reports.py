from datetime import datetime, date
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_excel_report(logs, student_name, date_from=None, date_to=None):
    """
    Генерирует Excel отчет по питанию ученика
    """
    # Подготовка данных для Excel
    data = []
    for log in logs:
        data.append({
            'Дата': log.created_at.strftime('%d.%m.%Y'),
            'Время': log.created_at.strftime('%H:%M'),
            'Блюдо': log.name,
            'Калории': log.calories,
            'Белки': log.protein,
            'Жиры': log.fat,
            'Углеводы': log.carbs
        })
    
    df = pd.DataFrame(data)
    
    # Создаем буфер в памяти для Excel файла
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Питание', index=False)
        worksheet = writer.sheets['Питание']
        
        # Форматирование
        for column in worksheet.columns:
            max_length = 0
            column = [cell for cell in column]
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = (max_length + 2)
            worksheet.column_dimensions[column[0].column_letter].width = adjusted_width

    output.seek(0)
    return output

def generate_word_report(logs, student_name, date_from=None, date_to=None):
    """
    Генерирует Word отчет по питанию ученика
    """
    doc = Document()
    
    # Заголовок
    title = doc.add_paragraph()
    title_run = title.add_run(f'Отчет по питанию ученика: {student_name}')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Период
    if date_from and date_to:
        period = doc.add_paragraph()
        period_run = period.add_run(f'Период: с {date_from.strftime("%d.%m.%Y")} по {date_to.strftime("%d.%m.%Y")}')
        period_run.font.size = Pt(12)
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Таблица с данными
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Дата', 'Время', 'Блюдо', 'Калории', 'Белки', 'Жиры', 'Углеводы']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Добавляем данные
    for log in logs:
        row_cells = table.add_row().cells
        row_cells[0].text = log.created_at.strftime('%d.%m.%Y')
        row_cells[1].text = log.created_at.strftime('%H:%M')
        row_cells[2].text = log.name
        row_cells[3].text = str(log.calories)
        row_cells[4].text = str(log.protein)
        row_cells[5].text = str(log.fat)
        row_cells[6].text = str(log.carbs)
    
    # Сохраняем в буфер
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output